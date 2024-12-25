from fastapi import FastAPI, BackgroundTasks, Request
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from backend.services.recorder import Recorder
from backend.services.summarizer import generate_summary
from docx import Document
from datetime import datetime
from backend.services.database import save_meeting_data
import cloudinary
import cloudinary.uploader
import cloudinary.api
from backend.services.utils.logging import logger
from backend.services.utils.exception import handle_exception

app = FastAPI()


cloudinary.config(
    cloud_name='df3o0herw',
    api_key='379621268338952',
    api_secret='vFYozAcP9RBBaHUiRmTBxU4GNWc'
)

app.mount("/template", StaticFiles(directory="template"), name="template")

transcription_file = "artifacts/generated_transcription.txt"
recorder = Recorder(transcription_file)

@app.post("/start_recording")
async def start_recording():
    try:
        recorder.start_recording()
        logger.info("Recording started.")
        return {"message": "Recording started."}
    except Exception as e:
        handle_exception(e, logger)
        return {"error": str(e)}

@app.post("/stop_recording")
async def stop_recording(request: Request, background_tasks: BackgroundTasks):
    try:
        recorder.stop_recording()
        transcription = recorder.get_transcriptions()
        meeting_details = await request.json()
        background_tasks.add_task(generate_summary_task, meeting_details)
        logger.info("Generated Transcriptions and Summary!")
        return {"message": "Generated Transcriptions and Summary!!", "transcription": transcription}
    except Exception as e:
        handle_exception(e, logger)
        return {"error": str(e)}

async def generate_summary_task(meeting_details: dict):
    try:
        with open(transcription_file, 'r') as file:
            text = file.read()
        
        result = generate_summary(text)
        agenda = result['agenda']
        summary = result['summary']
        resolution = result['resolution']
        
        # Update meeting details with generated summary data
        meeting_details.update({
            "date": meeting_details.get("date"),
            "time": meeting_details.get("time"),
            "initiator": meeting_details.get("initiator"),
            "verifier": meeting_details.get("verifier"),
            "participants": meeting_details.get("participants"),
            "transcription": text,
            "agenda": agenda,
            "summary": summary,
            "resolution": resolution,
            "timestamp": datetime.utcnow()
        })
        save_meeting_data(meeting_details)
            
        html_content = f"""
    <html>
    <head>
        <style>
            body {{
                font-family: Arial, sans-serif;
                background-color: #f9f9f9;
                margin: 0;
                padding: 20px;
                line-height: 1.6;
            }}
            .container {{
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: #fff;
                box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
                border-radius: 8px;
            }}
            h2 {{
                color: #333;
                font-size: 1.8em;
                margin-top: 20px;
                border-bottom: 2px solid #4c8caf;
                padding-bottom: 10px;
            }}
            p {{
                font-size: 1.1em;
                color: #555;
                background-color: #f0f0f0;
                padding: 10px;
                border-radius: 5px;
                margin-top: 10px;
            }}
            b {{
                color: #4c96af;
            }}
            @media screen and (max-width: 768px) {{
                .container {{
                    padding: 10px;
                }}
                h2 {{
                    font-size: 1.5em;
                }}
                p {{
                    font-size: 1em;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <h2><b>Agenda:</b></h2>
            <p>{agenda}</p>
            
            <h2><b>Resolution:</b></h2>
            <p>{resolution}</p>
            
            <h2><b>Summary:</b></h2>
            <p>{summary}</p>
        </div>
    </body>
    </html>
        """
        
        # Save the generated HTML content to a file
        with open("artifacts/generated_summary.html", 'w', encoding='utf-8') as summary_file:
            summary_file.write(html_content)

        print("Summary, agenda, and resolution generation complete.")
        print(agenda)
        print(summary)
        print(resolution)  

    except Exception as e:
        handle_exception(e, logger)

@app.post("/generate_document")
async def generate_document(request: Request):
    try:
        meeting_details = await request.json()
        with open(transcription_file, 'r') as file:
            transcription = file.read()

        result = generate_summary(transcription)
        agenda = result['agenda']
        summary = result['summary']
        resolution = result['resolution']

        doc = Document()
        doc.add_heading('Meeting Minutes', 0)
        doc.add_paragraph(f"Date: {meeting_details['date']}")
        doc.add_paragraph(f"Time: {meeting_details['time']}")
        doc.add_paragraph(f"Initiated by: {meeting_details['initiator']}")
        doc.add_paragraph(f"Minutes Verified by: {meeting_details['verifier']}")
        doc.add_paragraph(f"Number of Participants: {meeting_details['participants']}")

        doc.add_heading('Agenda', level=1)
        doc.add_paragraph(agenda)
        doc.add_heading('Resolution', level=1)
        doc.add_paragraph(resolution)
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)

        doc_path = "artifacts/meeting_minutes.docx"
        doc.save(doc_path)

        try:
            response = cloudinary.uploader.upload(doc_path, resource_type='raw')
            logger.info(f"File uploaded successfully! URL: {response['url']}")
            meeting_details.update({"document_url": response['url']})
            save_meeting_data(meeting_details)
        except Exception as e:
            handle_exception(e, logger)
        
        return JSONResponse(content={"message": "Document generated successfully.", "file_path": doc_path})

    except Exception as e:
        handle_exception(e, logger)
        return {"error": str(e)}

@app.get("/transcription")
async def get_transcription():
    try:
        return {"transcription": recorder.get_transcriptions()}
    except Exception as e:
        handle_exception(e, logger)
        return {"error": str(e)}

@app.get("/summary")
async def get_summary():
    try:
        with open("artifacts/generated_summary.html", 'r', encoding='utf-8') as summary_file:
            html_content = summary_file.read()
        return HTMLResponse(content=html_content, status_code=200)
    except FileNotFoundError:
        logger.warning("Summary file not found.")
        return HTMLResponse(content="<h2>No summary available yet.</h2>", status_code=404)

@app.get("/download_document")
async def download_document():
    try:
        doc_path = "artifacts/meeting_minutes.docx"
        return FileResponse(doc_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename="meeting_minutes.docx")
    except Exception as e:
        handle_exception(e, logger)
        return {"error": str(e)}



@app.get("/")
async def read_root():
    with open("template/index.html", 'r', encoding='utf-8') as f:
        content = f.read()
    return HTMLResponse(content=content, status_code=200)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
